"""
Content tools for Word Document Server.

These tools add various types of content to Word documents,
including headings, paragraphs, tables, images, and page breaks.
"""
import os
from typing import List, Optional, Dict, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension
from word_document_server.utils.document_utils import find_and_replace_text, insert_header_near_text, insert_numbered_list_near_text, insert_line_or_paragraph_near_text, replace_paragraph_block_below_header, replace_block_between_manual_anchors
from word_document_server.core.styles import ensure_heading_style, ensure_table_style


async def add_heading(filename: str, text: str, level: int = 1,
                      font_name: Optional[str] = None, font_size: Optional[int] = None,
                      bold: Optional[bool] = None, italic: Optional[bool] = None,
                      border_bottom: bool = False) -> str:
    """Add a heading to a Word document with optional formatting.

    Args:
        filename: Path to the Word document
        text: Heading text
        level: Heading level (1-9, where 1 is the highest level)
        font_name: Font family (e.g., 'Helvetica')
        font_size: Font size in points (e.g., 14)
        bold: True/False for bold text
        italic: True/False for italic text
        border_bottom: True to add bottom border (for section headers)
    """
    filename = ensure_docx_extension(filename)

    # Ensure level is converted to integer
    try:
        level = int(level)
    except (ValueError, TypeError):
        return "Invalid parameter: level must be an integer between 1 and 9"

    # Validate level range
    if level < 1 or level > 9:
        return f"Invalid heading level: {level}. Level must be between 1 and 9."

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."

    try:
        doc = Document(filename)

        # Ensure heading styles exist
        ensure_heading_style(doc)

        # Try to add heading with style
        try:
            heading = doc.add_heading(text, level=level)
        except Exception as style_error:
            # If style-based approach fails, use direct formatting
            heading = doc.add_paragraph(text)
            heading.style = doc.styles['Normal']
            if heading.runs:
                run = heading.runs[0]
                run.bold = True
                # Adjust size based on heading level
                if level == 1:
                    run.font.size = Pt(16)
                elif level == 2:
                    run.font.size = Pt(14)
                else:
                    run.font.size = Pt(12)

        # Apply formatting to all runs in the heading
        if any([font_name, font_size, bold is not None, italic is not None]):
            for run in heading.runs:
                if font_name:
                    run.font.name = font_name
                if font_size:
                    run.font.size = Pt(font_size)
                if bold is not None:
                    run.font.bold = bold
                if italic is not None:
                    run.font.italic = italic

        # Add bottom border if requested
        if border_bottom:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            pPr = heading._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')

            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')  # 0.5pt border
            bottom.set(qn('w:space'), '0')
            bottom.set(qn('w:color'), '000000')

            pBdr.append(bottom)
            pPr.append(pBdr)

        doc.save(filename)
        return f"Heading '{text}' (level {level}) added to {filename}"
    except Exception as e:
        return f"Failed to add heading: {str(e)}"


async def add_paragraph(filename: str, text: str, style: Optional[str] = None,
                        font_name: Optional[str] = None, font_size: Optional[int] = None,
                        bold: Optional[bool] = None, italic: Optional[bool] = None,
                        color: Optional[str] = None) -> str:
    """Add a paragraph to a Word document with optional formatting.

    Args:
        filename: Path to the Word document
        text: Paragraph text
        style: Optional paragraph style name
        font_name: Font family (e.g., 'Helvetica', 'Times New Roman')
        font_size: Font size in points (e.g., 14, 36)
        bold: True/False for bold text
        italic: True/False for italic text
        color: RGB color as hex string (e.g., '000000' for black)
    """
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."

    try:
        doc = Document(filename)
        paragraph = doc.add_paragraph(text)

        if style:
            try:
                paragraph.style = style
            except KeyError:
                # Style doesn't exist, use normal and report it
                paragraph.style = doc.styles['Normal']
                doc.save(filename)
                return f"Style '{style}' not found, paragraph added with default style to {filename}"

        # Apply formatting to all runs in the paragraph
        if any([font_name, font_size, bold is not None, italic is not None, color]):
            for run in paragraph.runs:
                if font_name:
                    run.font.name = font_name
                if font_size:
                    run.font.size = Pt(font_size)
                if bold is not None:
                    run.font.bold = bold
                if italic is not None:
                    run.font.italic = italic
                if color:
                    # Remove any '#' prefix if present
                    color_hex = color.lstrip('#')
                    run.font.color.rgb = RGBColor.from_string(color_hex)

        doc.save(filename)
        return f"Paragraph added to {filename}"
    except Exception as e:
        return f"Failed to add paragraph: {str(e)}"


async def add_table(filename: str, rows: int, cols: int, data: Optional[List[List[str]]] = None) -> str:
    """Add a table to a Word document.
    
    Args:
        filename: Path to the Word document
        rows: Number of rows in the table
        cols: Number of columns in the table
        data: Optional 2D array of data to fill the table
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."
    
    try:
        doc = Document(filename)
        table = doc.add_table(rows=rows, cols=cols)
        
        # Try to set the table style
        try:
            table.style = 'Table Grid'
        except KeyError:
            # If style doesn't exist, add basic borders
            pass
        
        # Fill table with data if provided
        if data:
            for i, row_data in enumerate(data):
                if i >= rows:
                    break
                for j, cell_text in enumerate(row_data):
                    if j >= cols:
                        break
                    table.cell(i, j).text = str(cell_text)
        
        doc.save(filename)
        return f"Table ({rows}x{cols}) added to {filename}"
    except Exception as e:
        return f"Failed to add table: {str(e)}"


async def add_picture(filename: str, image_path: str, width: Optional[float] = None) -> str:
    """Add an image to a Word document.
    
    Args:
        filename: Path to the Word document
        image_path: Path to the image file
        width: Optional width in inches (proportional scaling)
    """
    filename = ensure_docx_extension(filename)
    
    # Validate document existence
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Get absolute paths for better diagnostics
    abs_filename = os.path.abspath(filename)
    abs_image_path = os.path.abspath(image_path)
    
    # Validate image existence with improved error message
    if not os.path.exists(abs_image_path):
        return f"Image file not found: {abs_image_path}"
    
    # Check image file size
    try:
        image_size = os.path.getsize(abs_image_path) / 1024  # Size in KB
        if image_size <= 0:
            return f"Image file appears to be empty: {abs_image_path} (0 KB)"
    except Exception as size_error:
        return f"Error checking image file: {str(size_error)}"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(abs_filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."
    
    try:
        doc = Document(abs_filename)
        # Additional diagnostic info
        diagnostic = f"Attempting to add image ({abs_image_path}, {image_size:.2f} KB) to document ({abs_filename})"
        
        try:
            if width:
                doc.add_picture(abs_image_path, width=Inches(width))
            else:
                doc.add_picture(abs_image_path)
            doc.save(abs_filename)
            return f"Picture {image_path} added to {filename}"
        except Exception as inner_error:
            # More detailed error for the specific operation
            error_type = type(inner_error).__name__
            error_msg = str(inner_error)
            return f"Failed to add picture: {error_type} - {error_msg or 'No error details available'}\nDiagnostic info: {diagnostic}"
    except Exception as outer_error:
        # Fallback error handling
        error_type = type(outer_error).__name__
        error_msg = str(outer_error)
        return f"Document processing error: {error_type} - {error_msg or 'No error details available'}"


async def add_page_break(filename: str) -> str:
    """Add a page break to the document.
    
    Args:
        filename: Path to the Word document
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        doc.add_page_break()
        doc.save(filename)
        return f"Page break added to {filename}."
    except Exception as e:
        return f"Failed to add page break: {str(e)}"


async def add_table_of_contents(filename: str, title: str = "Table of Contents", max_level: int = 3) -> str:
    """Add a table of contents to a Word document based on heading styles.
    
    Args:
        filename: Path to the Word document
        title: Optional title for the table of contents
        max_level: Maximum heading level to include (1-9)
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        # Ensure max_level is within valid range
        max_level = max(1, min(max_level, 9))
        
        doc = Document(filename)
        
        # Collect headings and their positions
        headings = []
        for i, paragraph in enumerate(doc.paragraphs):
            # Check if paragraph style is a heading
            if paragraph.style and paragraph.style.name.startswith('Heading '):
                try:
                    # Extract heading level from style name
                    level = int(paragraph.style.name.split(' ')[1])
                    if level <= max_level:
                        headings.append({
                            'level': level,
                            'text': paragraph.text,
                            'position': i
                        })
                except (ValueError, IndexError):
                    # Skip if heading level can't be determined
                    pass
        
        if not headings:
            return f"No headings found in document {filename}. Table of contents not created."
        
        # Create a new document with the TOC
        toc_doc = Document()
        
        # Add title
        if title:
            toc_doc.add_heading(title, level=1)
        
        # Add TOC entries
        for heading in headings:
            # Indent based on level (using tab characters)
            indent = '    ' * (heading['level'] - 1)
            toc_doc.add_paragraph(f"{indent}{heading['text']}")
        
        # Add page break
        toc_doc.add_page_break()
        
        # Get content from original document
        for paragraph in doc.paragraphs:
            p = toc_doc.add_paragraph(paragraph.text)
            # Copy style if possible
            try:
                if paragraph.style:
                    p.style = paragraph.style.name
            except:
                pass
        
        # Copy tables
        for table in doc.tables:
            # Create a new table with the same dimensions
            new_table = toc_doc.add_table(rows=len(table.rows), cols=len(table.columns))
            # Copy cell contents
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    for paragraph in cell.paragraphs:
                        new_table.cell(i, j).text = paragraph.text
        
        # Save the new document with TOC
        toc_doc.save(filename)
        
        return f"Table of contents with {len(headings)} entries added to {filename}"
    except Exception as e:
        return f"Failed to add table of contents: {str(e)}"


async def delete_paragraph(filename: str, paragraph_index: int) -> str:
    """Delete a paragraph from a document.
    
    Args:
        filename: Path to the Word document
        paragraph_index: Index of the paragraph to delete (0-based)
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate paragraph index
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Invalid paragraph index. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1})."
        
        # Delete the paragraph (by removing its content and setting it empty)
        # Note: python-docx doesn't support true paragraph deletion, this is a workaround
        paragraph = doc.paragraphs[paragraph_index]
        p = paragraph._p
        p.getparent().remove(p)
        
        doc.save(filename)
        return f"Paragraph at index {paragraph_index} deleted successfully."
    except Exception as e:
        return f"Failed to delete paragraph: {str(e)}"


async def search_and_replace(filename: str, find_text: str, replace_text: str) -> str:
    """Search for text and replace all occurrences.
    
    Args:
        filename: Path to the Word document
        find_text: Text to search for
        replace_text: Text to replace with
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Perform find and replace
        count = find_and_replace_text(doc, find_text, replace_text)
        
        if count > 0:
            doc.save(filename)
            return f"Replaced {count} occurrence(s) of '{find_text}' with '{replace_text}'."
        else:
            return f"No occurrences of '{find_text}' found."
    except Exception as e:
        return f"Failed to search and replace: {str(e)}"

async def insert_header_near_text_tool(filename: str, target_text: str = None, header_title: str = "", position: str = 'after', header_style: str = 'Heading 1', target_paragraph_index: int = None) -> str:
    """Insert a header (with specified style) before or after the target paragraph. Specify by text or paragraph index."""
    return insert_header_near_text(filename, target_text, header_title, position, header_style, target_paragraph_index)

async def insert_numbered_list_near_text_tool(filename: str, target_text: str = None, list_items: list = None, position: str = 'after', target_paragraph_index: int = None, bullet_type: str = 'bullet') -> str:
    """Insert a bulleted or numbered list before or after the target paragraph. Specify by text or paragraph index."""
    return insert_numbered_list_near_text(filename, target_text, list_items, position, target_paragraph_index, bullet_type)

async def insert_line_or_paragraph_near_text_tool(filename: str, target_text: str = None, line_text: str = "", position: str = 'after', line_style: str = None, target_paragraph_index: int = None) -> str:
    """Insert a new line or paragraph (with specified or matched style) before or after the target paragraph. Specify by text or paragraph index."""
    return insert_line_or_paragraph_near_text(filename, target_text, line_text, position, line_style, target_paragraph_index)

async def replace_paragraph_block_below_header_tool(filename: str, header_text: str, new_paragraphs: list, detect_block_end_fn=None) -> str:
    """Reemplaza el bloque de párrafos debajo de un encabezado, evitando modificar TOC."""
    return replace_paragraph_block_below_header(filename, header_text, new_paragraphs, detect_block_end_fn)

async def replace_block_between_manual_anchors_tool(filename: str, start_anchor_text: str, new_paragraphs: list, end_anchor_text: str = None, match_fn=None, new_paragraph_style: str = None) -> str:
    """Replace all content between start_anchor_text and end_anchor_text (or next logical header if not provided)."""
    return replace_block_between_manual_anchors(filename, start_anchor_text, new_paragraphs, end_anchor_text, match_fn, new_paragraph_style)


async def add_row_to_table(filename: str, table_index: int, data: List[str],
                           row_index: Optional[int] = None) -> str:
    """Add a row to an existing table in a Word document.

    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        data: List of cell values for the new row (must match column count)
        row_index: Position to insert the row (0-based). None = append at the end.
    """
    filename = ensure_docx_extension(filename)

    # Ensure numeric parameters are the correct type
    try:
        table_index = int(table_index)
        if row_index is not None:
            row_index = int(row_index)
    except (ValueError, TypeError):
        return "Invalid parameter: table_index and row_index must be integers"

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."

    try:
        doc = Document(filename)

        # Validate table index
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Invalid table_index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1})."

        table = doc.tables[table_index]
        num_cols = len(table.columns)
        num_rows = len(table.rows)

        # Validate data length matches column count
        if len(data) != num_cols:
            return f"Data length ({len(data)}) does not match table column count ({num_cols})."

        # Validate row_index range
        if row_index is not None and (row_index < 0 or row_index > num_rows):
            return f"Invalid row_index ({row_index}). Valid range: 0-{num_rows} (0=before first row, {num_rows}=after last row)."

        if row_index is None or row_index == num_rows:
            # Append at the end
            new_row = table.add_row()
            for j, cell_text in enumerate(data):
                new_row.cells[j].text = str(cell_text)
            actual_index = num_rows
        else:
            # Insert at specific position using XML manipulation
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from copy import deepcopy

            # Create a new row by adding and then moving it
            new_row = table.add_row()
            for j, cell_text in enumerate(data):
                new_row.cells[j].text = str(cell_text)

            # Move the new row XML element to the desired position
            tbl = table._tbl
            tr_new = new_row._tr
            tbl.remove(tr_new)

            # Get the reference row to insert before
            tr_ref = table.rows[row_index]._tr
            tbl.insert(tbl.index(tr_ref), tr_new)
            actual_index = row_index

        doc.save(filename)
        new_row_count = len(table.rows)
        return f"Row added at index {actual_index} to table {table_index} in {filename} (now {new_row_count} rows x {num_cols} cols)"
    except Exception as e:
        return f"Failed to add row to table: {str(e)}"


async def add_column_to_table(filename: str, table_index: int, data: List[str],
                              col_index: Optional[int] = None) -> str:
    """Add a column to an existing table in a Word document.

    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        data: List of cell values for the new column (must match row count)
        col_index: Position to insert the column (0-based). None = append at the right end.
    """
    filename = ensure_docx_extension(filename)

    # Ensure numeric parameters are the correct type
    try:
        table_index = int(table_index)
        if col_index is not None:
            col_index = int(col_index)
    except (ValueError, TypeError):
        return "Invalid parameter: table_index and col_index must be integers"

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."

    try:
        doc = Document(filename)

        # Validate table index
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Invalid table_index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1})."

        table = doc.tables[table_index]
        num_rows = len(table.rows)
        num_cols = len(table.columns)

        # Validate data length matches row count
        if len(data) != num_rows:
            return f"Data length ({len(data)}) does not match table row count ({num_rows})."

        # Validate col_index range
        if col_index is not None and (col_index < 0 or col_index > num_cols):
            return f"Invalid col_index ({col_index}). Valid range: 0-{num_cols} (0=before first column, {num_cols}=after last column)."

        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        if col_index is None or col_index == num_cols:
            # Append at the right end
            table.add_column(Inches(1))  # Add column with default width
            for i, row in enumerate(table.rows):
                row.cells[num_cols].text = str(data[i])
            actual_index = num_cols
        else:
            # Insert at specific position using XML manipulation
            for i, row in enumerate(table.rows):
                tr = row._tr
                # Create new cell element
                new_tc = OxmlElement('w:tc')
                # Add paragraph with text
                p = OxmlElement('w:p')
                r = OxmlElement('w:r')
                t = OxmlElement('w:t')
                t.text = str(data[i])
                r.append(t)
                p.append(r)
                new_tc.append(p)

                # Insert before the reference cell
                ref_tc = row.cells[col_index]._tc
                tr.insert(tr.index(ref_tc), new_tc)
            actual_index = col_index

        doc.save(filename)
        new_col_count = len(table.columns)
        return f"Column added at index {actual_index} to table {table_index} in {filename} (now {num_rows} rows x {new_col_count} cols)"
    except Exception as e:
        return f"Failed to add column to table: {str(e)}"


async def merge_table_row_cells(filename: str, table_index: int, row_index: int,
                                start_col_index: int, end_col_index: int,
                                text: Optional[str] = None) -> str:
    """Merge cells in the same row of an existing table.

    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        row_index: Index of the row containing the cells (0-based)
        start_col_index: Starting column index to merge (0-based)
        end_col_index: Ending column index to merge (0-based)
        text: Optional text to set in the merged cell. If None, contents are concatenated.
    """
    filename = ensure_docx_extension(filename)

    try:
        table_index = int(table_index)
        row_index = int(row_index)
        start_col_index = int(start_col_index)
        end_col_index = int(end_col_index)
    except (ValueError, TypeError):
        return "Invalid parameter: indices must be integers"

    if start_col_index >= end_col_index:
        return "Invalid span: start_col_index must be less than end_col_index"

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."

    try:
        doc = Document(filename)

        if table_index < 0 or table_index >= len(doc.tables):
            return f"Invalid table_index. Document has {len(doc.tables)} tables."

        table = doc.tables[table_index]
        num_rows = len(table.rows)
        num_cols = len(table.columns)

        if row_index < 0 or row_index >= num_rows:
            return f"Invalid row_index ({row_index}). Valid range: 0-{num_rows-1}."

        if start_col_index < 0 or end_col_index >= num_cols:
            return f"Invalid column span ({start_col_index} to {end_col_index}). Valid range: 0-{num_cols-1}."

        cell1 = table.cell(row_index, start_col_index)
        cell2 = table.cell(row_index, end_col_index)
        
        merged_cell = cell1.merge(cell2)

        if text is not None:
            merged_cell.text = str(text)

        doc.save(filename)
        return f"Merged row {row_index} cells from col {start_col_index} to {end_col_index} in table {table_index}"
    except Exception as e:
        return f"Failed to merge row cells: {str(e)}"


async def merge_table_column_cells(filename: str, table_index: int, col_index: int,
                                   start_row_index: int, end_row_index: int,
                                   text: Optional[str] = None) -> str:
    """Merge cells in the same column of an existing table.

    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        col_index: Index of the column containing the cells (0-based)
        start_row_index: Starting row index to merge (0-based)
        end_row_index: Ending row index to merge (0-based)
        text: Optional text to set in the merged cell. If None, contents are concatenated.
    """
    filename = ensure_docx_extension(filename)

    try:
        table_index = int(table_index)
        col_index = int(col_index)
        start_row_index = int(start_row_index)
        end_row_index = int(end_row_index)
    except (ValueError, TypeError):
        return "Invalid parameter: indices must be integers"

    if start_row_index >= end_row_index:
        return "Invalid span: start_row_index must be less than end_row_index"

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."

    try:
        doc = Document(filename)

        if table_index < 0 or table_index >= len(doc.tables):
            return f"Invalid table_index. Document has {len(doc.tables)} tables."

        table = doc.tables[table_index]
        num_rows = len(table.rows)
        num_cols = len(table.columns)

        if col_index < 0 or col_index >= num_cols:
            return f"Invalid col_index ({col_index}). Valid range: 0-{num_cols-1}."

        if start_row_index < 0 or end_row_index >= num_rows:
            return f"Invalid row span ({start_row_index} to {end_row_index}). Valid range: 0-{num_rows-1}."

        cell1 = table.cell(start_row_index, col_index)
        cell2 = table.cell(end_row_index, col_index)
        
        merged_cell = cell1.merge(cell2)

        if text is not None:
            merged_cell.text = str(text)

        doc.save(filename)
        return f"Merged col {col_index} cells from row {start_row_index} to {end_row_index} in table {table_index}"
    except Exception as e:
        return f"Failed to merge column cells: {str(e)}"


async def remove_row_from_table(filename: str, table_index: int, row_index: int) -> str:
    """Remove a row from an existing table in a Word document.

    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        row_index: Index of the row to remove (0-based)
    """
    filename = ensure_docx_extension(filename)

    try:
        table_index = int(table_index)
        row_index = int(row_index)
    except (ValueError, TypeError):
        return "Invalid parameter: indices must be integers"

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."

    try:
        doc = Document(filename)

        if table_index < 0 or table_index >= len(doc.tables):
            return f"Invalid table_index. Document has {len(doc.tables)} tables."

        table = doc.tables[table_index]
        num_rows = len(table.rows)

        if row_index < 0 or row_index >= num_rows:
            return f"Invalid row_index ({row_index}). Valid range: 0-{num_rows-1}."

        tr = table.rows[row_index]._tr
        table._tbl.remove(tr)

        doc.save(filename)
        return f"Removed row {row_index} from table {table_index}. Remaining rows: {num_rows - 1}"
    except Exception as e:
        return f"Failed to remove row from table: {str(e)}"


async def remove_column_from_table(filename: str, table_index: int, col_index: int) -> str:
    """Remove a column from an existing table in a Word document.

    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        col_index: Index of the column to remove (0-based)
    """
    filename = ensure_docx_extension(filename)

    try:
        table_index = int(table_index)
        col_index = int(col_index)
    except (ValueError, TypeError):
        return "Invalid parameter: indices must be integers"

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."

    try:
        doc = Document(filename)
        from docx.oxml.ns import qn

        if table_index < 0 or table_index >= len(doc.tables):
            return f"Invalid table_index. Document has {len(doc.tables)} tables."

        table = doc.tables[table_index]
        num_cols = len(table.columns)

        if col_index < 0 or col_index >= num_cols:
            return f"Invalid col_index ({col_index}). Valid range: 0-{num_cols-1}."

        for row in table.rows:
            cells = row.cells
            if col_index < len(cells):
                tc = cells[col_index]._tc
                try:
                    row._tr.remove(tc)
                except ValueError:
                    # In case of horizontal spans, this _tc might have been removed already in a previous row check
                    pass

        # Cleanup <w:gridCol> if exists
        tblGrid = table._tbl.tblGrid
        if tblGrid is not None:
            gridCols = tblGrid.findall(qn('w:gridCol'))
            if len(gridCols) > col_index:
                tblGrid.remove(gridCols[col_index])

        doc.save(filename)
        return f"Removed column {col_index} from table {table_index}. Remaining cols: {num_cols - 1}"
    except Exception as e:
        return f"Failed to remove column from table: {str(e)}"


async def unmerge_all_table_cells(filename: str, table_index: int) -> str:
    """Unmerge all horizontally and vertically merged cells in an existing table.

    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
    """
    filename = ensure_docx_extension(filename)

    try:
        table_index = int(table_index)
    except (ValueError, TypeError):
        return "Invalid parameter: indices must be integers"

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."

    try:
        doc = Document(filename)
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        if table_index < 0 or table_index >= len(doc.tables):
            return f"Invalid table_index. Document has {len(doc.tables)} tables."

        table = doc.tables[table_index]
        changes_made = 0

        for row in table.rows:
            tr = row._tr
            tcs = tr.findall(qn('w:tc'))
            
            for tc in tcs:
                tcPr = tc.tcPr
                if tcPr is not None:
                    # Fix horizontal merges (gridSpan)
                    gridSpan = tcPr.gridSpan
                    if gridSpan is not None:
                        span_val = gridSpan.val
                        tcPr.remove(gridSpan)
                        changes_made += 1
                        # Create empty cells to restore the row grid
                        for _ in range(span_val - 1):
                            new_tc = OxmlElement('w:tc')
                            new_tc.append(OxmlElement('w:p'))
                            tc.addnext(new_tc)
                    
                    # Fix vertical merges (vMerge)
                    vMerge = tcPr.vMerge
                    if vMerge is not None:
                        tcPr.remove(vMerge)
                        changes_made += 1

        doc.save(filename)
        return f"Unmerged all cells in table {table_index}. Processed {changes_made} merge tags."
    except Exception as e:
        return f"Failed to unmerge table cells: {str(e)}"
